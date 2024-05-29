import PyPDF2
import fitz  # PyMuPDF
from docx import Document
from PIL import Image
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def pdf_to_image(pdf_path, image_path):
    pdf_document = fitz.open(pdf_path)
    for page_number in range(len(pdf_document)):
        page = pdf_document[page_number]
        image = page.get_pixmap()
        image.save(f"{image_path}_page_{page_number + 1}.png")

def pdf_to_text(pdf_path, text_path):
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page_number in range(len(reader.pages)):
            text += reader.pages[page_number].extract_text()

        with open(text_path, 'w', encoding='utf-8') as text_file:
            text_file.write(text)

def text_to_document(text_path, doc_path):
    document = Document()
    with open(text_path, 'r', encoding='utf-8') as text_file:
        for line in text_file:
            paragraph = document.add_paragraph(line.strip())
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = paragraph.runs[0]
            run.font.size = Pt(12)  # Set font size to 12pt (adjust as needed)
            # You can add more formatting options here

    document.save(doc_path)

# Example usage
pdf_file = r"C:\Users\DELL\Downloads\kavya junuthula (3).pdf"
image_output_path =r"C:\Users\DELL\Downloads"
text_output_path=r"C:\Users\DELL\textfile.txt"
doc_output_path =r"C:\Users\DELL\document.docx"

pdf_to_image(pdf_file, image_output_path)
pdf_to_text(pdf_file, text_output_path)
text_to_document(text_output_path, doc_output_path)
Footer
© 2024 GitHub, Inc.
Footer navigation
Terms
Privacy
Security
Status
import PyPDF2
import fitz  # PyMuPDF
from docx import Document
from PIL import Image
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def pdf_to_image(pdf_path, image_path):
    pdf_document = fitz.open(pdf_path)
    for page_number in range(len(pdf_document)):
        page = pdf_document[page_number]
        image = page.get_pixmap()
        image.save(f"{image_path}_page_{page_number + 1}.png")

def pdf_to_text(pdf_path, text_path):
    with open(pdf_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        text = ''
        for page_number in range(len(reader.pages)):
            text += reader.pages[page_number].extract_text()

        with open(text_path, 'w', encoding='utf-8') as text_file:
            text_file.write(text)

def text_to_document(text_path, doc_path):
    document = Document()
    with open(text_path, 'r', encoding='utf-8') as text_file:
        for line in text_file:
            paragraph = document.add_paragraph(line.strip())
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            run = paragraph.runs[0]
            run.font.size = Pt(12)  # Set font size to 12pt (adjust as needed)
            # You can add more formatting options here

    document.save(doc_path)

# Example usage
pdf_file = r"C:\Users\DELL\Downloads\Deepa Ajmeera (3).pdf"
image_output_path =r"C:\Users\DELL\Downloads"
text_output_path=r"C:\Users\DELL\textfile.txt"
doc_output_path =r"C:\Users\DELL\document.docx"

pdf_to_image(pdf_file, image_output_path)
pdf_to_text(pdf_file, text_output_path)
text_to_document(text_output_path, doc_output_path)
